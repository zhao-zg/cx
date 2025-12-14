# -*- coding: utf-8 -*-
"""
Word文档解析器
"""
import re
from docx import Document
from typing import List, Optional
from .models import Chapter, Content, TrainingData


class WordParser:
    """Word文档解析器"""
    
    # 样式映射
    STYLE_MAP = {
        '121文章篇题': 'chapter_title',
        '131文章大点': 'section_level1',  # 壹贰叁
        '132文章中点': 'section_level2',  # 一二三
        '133文章小点': 'section_level3',  # 1 2 3
        '134文章小a点': 'section_level4',  # a b c
        '8888文章正文': 'content'
    }
    
    # 大点层级标识（壹贰叁...）
    LEVEL1_MARKERS = ['壹', '贰', '叁', '肆', '伍', '陆', '柒', '捌', '玖', '拾']
    
    def __init__(self):
        self.current_chapter = None
        self.current_level1 = None
        self.current_level2 = None
        self.current_level3 = None
    
    def parse_listen_doc(self, docx_path: str) -> List[Chapter]:
        """
        解析听抄文档
        
        Args:
            docx_path: Word文档路径
            
        Returns:
            章节列表
        """
        doc = Document(docx_path)
        chapters = []
        
        for para in doc.paragraphs:
            style_type = self.STYLE_MAP.get(para.style.name)
            text = para.text.strip()
            
            if not text:
                continue
            
            if style_type == 'chapter_title':
                # 新的篇章
                if self.current_chapter:
                    chapters.append(self.current_chapter)
                
                chapter_num = self._extract_chapter_number(text)
                self.current_chapter = Chapter(
                    number=chapter_num,
                    title=text
                )
                self.current_level1 = None
                self.current_level2 = None
                self.current_level3 = None
                
            elif style_type == 'section_level1':
                # 大点（壹贰叁）
                if self.current_chapter:
                    level = self._extract_level_marker(text)
                    self.current_level1 = Content(level=level, title=text)
                    self.current_chapter.add_section(self.current_level1)
                    self.current_level2 = None
                    self.current_level3 = None
                    
            elif style_type == 'section_level2':
                # 中点（一二三）
                if self.current_level1:
                    level = self._extract_level_marker(text)
                    self.current_level2 = Content(level=level, title=text)
                    self.current_level1.add_child(self.current_level2)
                    self.current_level3 = None
                    
            elif style_type == 'section_level3':
                # 小点（1 2 3）
                if self.current_level2:
                    level = self._extract_level_marker(text)
                    self.current_level3 = Content(level=level, title=text)
                    self.current_level2.add_child(self.current_level3)
                    
            elif style_type == 'section_level4':
                # 小a点（a b c）
                if self.current_level3:
                    level = self._extract_level_marker(text)
                    level4 = Content(level=level, title=text)
                    self.current_level3.add_child(level4)
                    
            elif style_type == 'content':
                # 正文内容
                if self.current_level3:
                    self.current_level3.add_content(text)
                elif self.current_level2:
                    self.current_level2.add_content(text)
                elif self.current_level1:
                    self.current_level1.add_content(text)
                elif self.current_chapter:
                    self.current_chapter.listen_content.append(text)
        
        # 添加最后一章
        if self.current_chapter:
            chapters.append(self.current_chapter)
        
        return chapters
    
    def _extract_chapter_number(self, text: str) -> int:
        """从标题中提取篇章编号"""
        match = re.search(r'第([一二三四五六七八九])篇', text)
        if match:
            chinese_num = match.group(1)
            num_map = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
                      '六': 6, '七': 7, '八': 8, '九': 9}
            return num_map.get(chinese_num, 1)
        return 1
    
    def _extract_level_marker(self, text: str) -> str:
        """提取层级标识（壹、一、1、a等）"""
        # 尝试匹配各种层级标识
        patterns = [
            r'^([壹贰叁肆伍陆柒捌玖拾]+)\s',  # 壹贰叁...
            r'^([一二三四五六七八九十]+)\s',  # 一二三...
            r'^(\d+)\s',                      # 1 2 3...
            r'^([a-z])\s',                    # a b c...
        ]
        
        for pattern in patterns:
            match = re.match(pattern, text)
            if match:
                return match.group(1)
        
        return ''
    
    def parse_scripture_doc(self, docx_path: str, chapters: List[Chapter]):
        """
        解析经文文档，补充经文信息到已有章节
        
        Args:
            docx_path: Word文档路径
            chapters: 已解析的章节列表
        """
        doc = Document(docx_path)
        current_chapter_num = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 查找篇章编号
            match = re.search(r'第([一二三四五六七八九])篇', text)
            if match:
                chinese_num = match.group(1)
                num_map = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
                          '六': 6, '七': 7, '八': 8, '九': 9}
                current_chapter_num = num_map.get(chinese_num, 0)
            
            # 提取经文引用（如：太五3、8）
            if current_chapter_num > 0 and '—' in text:
                scripture = text.split('—')[-1].strip()
                for chapter in chapters:
                    if chapter.number == current_chapter_num:
                        chapter.scripture = scripture
                        break


def parse_training_docs(listen_path: str, scripture_path: str, 
                       title: str, subtitle: str, year: int, season: str) -> TrainingData:
    """
    解析训练文档，生成完整的训练数据
    
    Args:
        listen_path: 听抄文档路径
        scripture_path: 经文文档路径
        title: 训练标题
        subtitle: 副标题
        year: 年份
        season: 季节
        
    Returns:
        训练数据对象
    """
    parser = WordParser()
    
    # 解析听抄文档
    chapters = parser.parse_listen_doc(listen_path)
    
    # 解析经文文档并补充信息
    parser.parse_scripture_doc(scripture_path, chapters)
    
    # 创建训练数据对象
    training_data = TrainingData(
        title=title,
        subtitle=subtitle,
        year=year,
        season=season
    )
    
    for chapter in chapters:
        training_data.add_chapter(chapter)
    
    return training_data
