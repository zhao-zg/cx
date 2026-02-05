# -*- coding: utf-8 -*-
"""
改进的Word文档解析器 - 正确区分数据来源
"""
import re
import os
import shutil
import subprocess
from docx import Document
from typing import List, Optional
from .models import Chapter, Content, TrainingData, MorningRevival


def load_document(doc_path: str):
    """
    加载Word文档,自动识别 .doc 和 .docx 格式
    
    Args:
        doc_path: 文档路径
    
    Returns:
        Document 对象或 None
    """
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"文档不存在: {doc_path}")
    
    # 根据扩展名判断格式
    ext = os.path.splitext(doc_path)[1].lower()
    
    if ext == '.docx':
        # 使用 python-docx 解析 .docx
        return Document(doc_path)
    elif ext == '.doc':
        # .doc格式需要先转换为.docx
        # 尝试使用LibreOffice进行转换（跨平台方案）
        try:
            import tempfile
            
            # 创建临时目录
            temp_dir = tempfile.mkdtemp()
            abs_path = os.path.abspath(doc_path)
            
            # 尝试找到LibreOffice/soffice命令
            soffice_commands = [
                'soffice',           # Linux/Mac
                'libreoffice',       # Linux
                r'C:\Program Files\LibreOffice\program\soffice.exe',  # Windows
                r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
            ]
            
            soffice_path = None
            for cmd in soffice_commands:
                if shutil.which(cmd) or os.path.exists(cmd):
                    soffice_path = cmd
                    break
            
            if soffice_path:
                print(f"    ⏳ 正在转换 .doc 文件...")
                # 使用LibreOffice转换
                result = subprocess.run(
                    [soffice_path, '--headless', '--convert-to', 'docx', '--outdir', temp_dir, abs_path],
                    capture_output=True,
                    timeout=60
                )
                
                if result.returncode == 0:
                    # 查找转换后的文件
                    docx_name = os.path.splitext(os.path.basename(doc_path))[0] + '.docx'
                    temp_docx = os.path.join(temp_dir, docx_name)
                    
                    if os.path.exists(temp_docx):
                        print(f"    ✓ 转换成功，继续处理...")
                        doc = Document(temp_docx)
                        # 清理临时文件
                        shutil.rmtree(temp_dir, ignore_errors=True)
                        return doc
            
            # LibreOffice不可用或转换失败
            shutil.rmtree(temp_dir, ignore_errors=True)
            
            # 提供友好的错误提示
            print("\n" + "="*60)
            print("⚠ 无法自动转换 .doc 文件")
            print("="*60)
            print("\n请选择以下解决方案之一：")
            print("\n方案 1: 手动转换（最快）")
            print(f"  1. 在 Word 中打开: {doc_path}")
            print("  2. 另存为 .docx 格式")
            print("  3. 重新运行此程序")
            print("\n方案 2: 安装 LibreOffice（自动化）")
            print("  运行转换工具: python convert_doc_to_docx.py")
            print("  工具会自动检测系统并引导安装")
            print("\n方案 3: 使用在线转换")
            print("  https://www.online-convert.com/")
            print("  https://www.zamzar.com/")
            print("\n" + "="*60)
            
            raise ImportError(
                f"无法转换 .doc 文件: {os.path.basename(doc_path)}\n"
                "请安装 LibreOffice 或手动转换为 .docx 格式"
            )
            
        except subprocess.TimeoutExpired:
            raise Exception("LibreOffice 转换超时（60秒）")
        except Exception as e:
            if isinstance(e, ImportError):
                raise
            raise Exception(f"解析 .doc 文件失败: {e}")
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


class ImprovedParser:
    """改进的解析器"""
    
    STYLE_MAP = {
        # 秋季 .docx 样式
        '121文章篇题': 'chapter_title',
        '131文章大点': 'section_level1',
        '132文章中点': 'section_level2',
        '133文章小点': 'section_level3',
        '134文章小a点': 'section_level4',
        '8888文章正文': 'content',
        # 夏季 .doc 样式
        '０ａ總題': 'chapter_title',
        '職事信息大標': 'section_level1',
        '職事信息中標': 'section_level2',
        '职事小标题': 'section_level3',
        '信息正文18': 'content',
        '信息正文3': 'content',
        '信息正文17': 'content',
        '職事信息': 'content'
    }
    
    # 预编译正则表达式
    WEEK_OUTLINE_PATTERN = re.compile(r'^第([一二三四五六七八九十]+)周[　\s]*•[　\s]*纲目')
    DAY_PATTERN = re.compile(r'^第([一二三四五六七八九十]+)周[　\s]*•[　\s]*周([一二三四五六七])')
    LEVEL1_PATTERN = re.compile(r'^([壹贰叁肆伍陆柒捌玖拾])[　\s]+(.*)')
    LEVEL2_PATTERN = re.compile(r'^([一二三四五六七八九十百]+)[　\s]+(.*)')
    LEVEL3_PATTERN = re.compile(r'^(\d+)[　\s]+(.*)')
    # 经文格式：太5:3	经文内容... 或 腓2:5	经文内容... 或 太五3	经文内容...
    # 支持两种格式: 1) 书卷+中文数字+阿拉伯数字 (太五3), 2) 书卷+阿拉伯数字 (腓2:5)
    VERSE_PATTERN = re.compile(r'^([创出利民申书士得撒王代拉尼斯伯诗箴传歌赛耶哀结但何珥摩俄拿弥鸿哈番该亚玛太可路约徒罗林加弗腓西帖提门多彼约犹启](?:[一二三四五六七八九十后前]\d+|\d+):\d+)[　\s\t]+(.+)')
    
    def __init__(self, output_dir: str = 'output'):
        self.output_dir = output_dir
        self.training_title = ""  # 训练主标题
        self.training_subtitle = ""  # 训练副标题
        self.first_week_number = None  # 记录第一个晨兴文档的起始周数
        self.reset_state()
    
    def reset_state(self):
        """重置解析状态"""
        self.current_chapter = None
        self.current_level1 = None
        self.current_level2 = None
        self.current_level3 = None
        self.verse_cache = {}  # 缓存已出现的经文范围内容
    
    def _chinese_to_number(self, chinese_str: str) -> Optional[int]:
        """
        转换中文数字为阿拉伯数字
        例如: "三十一" -> 31, "十二" -> 12, "七" -> 7
        """
        if not chinese_str:
            return None
        
        chinese_numerals = {
            '零': 0, '一': 1, '二': 2, '三': 3, '四': 4,
            '五': 5, '六': 6, '七': 7, '八': 8, '九': 9,
            '十': 10, '百': 100
        }
        
        try:
            # 特殊情况: "十" 单独出现时表示10
            if chinese_str == '十':
                return 10
            
            result = 0
            temp = 0
            
            for char in chinese_str:
                if char not in chinese_numerals:
                    return None
                
                num = chinese_numerals[char]
                
                if num >= 10:  # 十或百
                    if temp == 0:
                        temp = 1
                    result += temp * num
                    temp = 0
                else:
                    temp = num
            
            result += temp
            return result if result > 0 else None
        except:
            return None
    
    def _is_verse_line(self, text: str) -> bool:
        """
        检测文本是否是经文格式
        格式：太5:3	经文内容... 或 腓2:5	经文内容...
        """
        if not text:
            return False
        return bool(self.VERSE_PATTERN.match(text))
    
    def _extract_verse_range(self, text: str) -> tuple:
        """
        提取经文范围信息
        例如: "腓2:5~11 从略。" -> ('腓2', 5, 11, True)
              "腓2:5" -> ('腓2', 5, 5, False)
        返回: (book, start_verse, end_verse, is_omitted)
        """
        # 匹配 "腓2:5~11 从略。" 或 "腓2:5~11" - 支持两种格式
        range_match = re.match(r'^([创出利民申书士得撒王代拉尼斯伯诗箴传歌赛耶哀结但何珥摩俄拿弥鸿哈番该亚玛太可路约徒罗林加弗腓西帖提多彼约犹启](?:[一二三四五六七八九十后前]\d+|\d+)):(\d+)~(\d+)', text)
        if range_match:
            book = range_match.group(1)
            start = int(range_match.group(2))
            end = int(range_match.group(3))
            is_omitted = '从略' in text
            return (book, start, end, is_omitted)
        
        # 匹配单节 "腓2:5"
        single_match = re.match(r'^([创出利民申书士得撒王代拉尼斯伯诗箴传歌赛耶哀结但何珥摩俄拿弥鸿哈番该亚玛太可路约徒罗林加弗腓西帖提多彼约犹启](?:[一二三四五六七八九十后前]\d+|\d+)):(\d+)', text)
        if single_match:
            book = single_match.group(1)
            verse = int(single_match.group(2))
            return (book, verse, verse, False)
        
        return None
    
    def _get_verse_range_key(self, book: str, start: int, end: int) -> str:
        """生成经文范围的缓存键"""
        return f"{book}:{start}~{end}"
    
    def _cache_verse(self, text: str):
        """
        缓存单节经文到verse_cache
        例如: "腓2:5	你们里面要思念..." -> cache['腓2:5'] = '腓2:5	你们里面要思念...'
        """
        match = self.VERSE_PATTERN.match(text)
        if match:
            verse_ref = match.group(1)  # 例如: 腓2:5
            self.verse_cache[verse_ref] = text
    
    def _get_cached_verse_range(self, book: str, start: int, end: int) -> str:
        """
        从缓存中获取经文范围的内容
        例如: ('腓2', 5, 11) -> 返回 腓2:5 到 腓2:11 的所有经文
        """
        verses = []
        for verse_num in range(start, end + 1):
            verse_key = f"{book}:{verse_num}"
            if verse_key in self.verse_cache:
                verses.append(self.verse_cache[verse_key])
        return '\n'.join(verses) if verses else ''
    
    def parse_outline_doc(self, docx_path: str) -> List[Chapter]:
        """
        解析纲目文档（经文.docx/.doc）- 提取大纲结构和职事信息摘录
        """
        # print(f"  解析纲目结构（{os.path.basename(docx_path)}）...")
        doc = load_document(docx_path)
        chapters = []
        self.reset_state()
        in_content_section = False
        chapter_title_buffer = ""  # 用于累积多行标题
        current_node = None  # 当前正在处理的节点
        collecting_verses = False  # 是否正在收集绋文
        collecting_ministry = False  # 是否正在收集职事信息摘录
        ministry_buffer = []  # 职事信息摘录缓冲区
        
        # 首先在文档开头提取标题信息
        title_parts = []
        subtitle_found = False
        for i, para in enumerate(doc.paragraphs[:15]):  # 只检查前15段
            text = para.text.strip()
            if not text:
                continue
                
            # 停止条件：遇到"目录"或"第X篇"
            if re.match(r'目\s*录|^第[一二三四五六七八九十]+篇', text):
                break
                
            # 识别副标题（"总题："后面的内容）
            if text == "总题：" and i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1].text.strip()
                if next_para and not re.match(r'目\s*录|^第[一二三四五六七八九十]+篇', next_para):
                    self.training_subtitle = next_para
                    print(f"  ✓ 识别副标题: {self.training_subtitle}")
                    subtitle_found = True
                continue
                
            # 特殊处理：如果第一行包含常见的训练主题内容，直接作为副标题
            if i == 0 and not subtitle_found and len(text) > 5 and len(text) < 50:
                if re.search(r'经历|享受|彰显|基督|神|生命|召会', text):
                    self.training_subtitle = text
                    print(f"  ✓ 识别副标题（首行）: {self.training_subtitle}")
                    subtitle_found = True
                    continue
                
            # 收集可能的标题部分
            if not subtitle_found and text and len(text) < 50:  # 标题通常不会太长
                # 排除一些明显不是标题的内容
                if not re.match(r'^\(.+\)$|^\d+$|^页\s*\d+', text):
                    title_parts.append(text)
        
        # 组合主标题
        if title_parts:
            # 尝试找到包含"训练"或"特会"的组合
            for i in range(len(title_parts)):
                for j in range(i + 1, len(title_parts) + 1):
                    candidate_title = " ".join(title_parts[i:j])
                    if re.search(r'训练|特会', candidate_title):
                        self.training_title = candidate_title
                        print(f"  ✓ 识别标题: {self.training_title}")
                        break
                if self.training_title:
                    break
            
            # 如果没有找到包含"训练"或"特会"的组合，使用包含时间的组合
            if not self.training_title:
                for i in range(len(title_parts)):
                    for j in range(i + 1, len(title_parts) + 1):
                        candidate_title = " ".join(title_parts[i:j])
                        if re.search(r'[二〇一三四五六七八九零]+年|20\d+年', candidate_title):
                            self.training_title = candidate_title
                            print(f"  ✓ 识别标题（按时间）: {self.training_title}")
                            break
                    if self.training_title:
                        break
        
        # 如果还没有找到标题，尝试从文件路径推断
        if not self.training_title:
            # 支持常见的训练类型: 春季/夏季/秋季/冬季/感恩节/圣诞节等
            training_types = ['夏季', '秋季', '春季', '冬季', '感恩节', '圣诞节', '国殇节', '新年']
            path_lower = docx_path.lower()
            
            for training_type in training_types:
                if training_type in path_lower or training_type.lower() in path_lower:
                    self.training_title = f"{training_type}训练"
                    print(f"  ✓ 从路径推断标题: {self.training_title}")
                    break
            
            # 如果还是没找到,使用通用标题
            if not self.training_title:
                self.training_title = "训练"
                print(f"  ⚠ 无法识别训练类型，使用通用标题: {self.training_title}")

        for para in doc.paragraphs:
            text = para.text.strip()
            
            # 如果是verses样式或经文格式，添加到当前节点的scripture
            is_verse_by_style = para.style and (para.style.name == 'verses' or para.style.name == '０c 經節')
            is_verse_by_format = self._is_verse_line(text)
            
            if (is_verse_by_style or is_verse_by_format) and current_node and text:
                # 处理"从略"占位符：用缓存的经文范围替换
                if '从略' in text:
                    verse_info = self._extract_verse_range(text)
                    if verse_info:
                        book, start, end, is_omitted = verse_info
                        if is_omitted:
                            # 从缓存获取该范围的经文
                            cached_verses = self._get_cached_verse_range(book, start, end)
                            if cached_verses:
                                if current_node.scripture:
                                    current_node.scripture += '\n' + cached_verses
                                else:
                                    current_node.scripture = cached_verses
                    continue
                
                # 正常经文：缓存并添加
                self._cache_verse(text)
                if current_node.scripture:
                    current_node.scripture += '\n' + text
                else:
                    current_node.scripture = text
                continue
            
            if not text:
                continue
            
            # 查找第一个实际内容的开始（跳过目录）
            if not in_content_section:
                # 当找到一个单独的"第X篇"（目录后的正文开始）
                if re.match(r'^第[一二三四五六七八九十]+篇$', text):
                    in_content_section = True
                    chapter_title_buffer = text
                    continue
                continue
            
            # 处理篇章标题（支持第1-99篇）
            if re.match(r'^第[一二三四五六七八九十百]+篇', text):
                # 保存上一篇
                if self.current_chapter:
                    # 如果正在收集职事信息，先保存到当前章节
                    if collecting_ministry and ministry_buffer:
                        self.current_chapter.ministry_excerpt = '\n\n'.join(ministry_buffer)
                        # print(f"  [职事信息摘录] 保存第{self.current_chapter.number}篇职事信息摘录，共 {len(ministry_buffer)} 段")
                    chapters.append(self.current_chapter)
                    self.reset_state()
                    current_node = None  # 重置当前节点
                
                # 重置职事信息收集状态
                ministry_buffer = []
                collecting_ministry = False
                
                # 开始新篇
                chapter_title_buffer = text
                continue
            
                # 如果有标题缓冲，继续累积直到遇到其他内容
            if chapter_title_buffer:
                # 检查是否是标题的延续（不是"MC 诗歌"、"读经"、诗歌编号等）
                if not any([
                    text.startswith('MC '),
                    text.startswith('JL '),
                    text.startswith('SC '),
                    text.startswith('RM '),
                    text.startswith('NL/'),
                    text.startswith('HL '),
                    text.startswith('RA '),
                    text.startswith('读经：'),
                    '诗歌：' in text
                ]):
                    chapter_title_buffer += text
                    continue
                else:
                    # 标题结束，创建章节
                    chapter_num = self._extract_chapter_number(chapter_title_buffer)
                    
                    # 从标题中提取诗歌信息（夏季格式："...EM 诗歌: s250, s432"）
                    hymn_match = re.search(r'(EM|RK)\s*诗歌[：:]\s*([^，。]+)', chapter_title_buffer)
                    if hymn_match:
                        hymn_number = f"{hymn_match.group(1)} 诗歌: {hymn_match.group(2)}"
                    else:
                        hymn_number = ""
                    
                    # 从完整标题中提取实际标题（去除"第X篇"部分和诗歌信息）
                    clean_title = re.sub(r'^第[一二三四五六七八九十百]+篇\s*', '', chapter_title_buffer)
                    clean_title = re.sub(r'(EM|RK)\s*诗歌[：:]\s*[^，。]*', '', clean_title).strip()
                    
                    self.current_chapter = Chapter(
                        number=chapter_num,
                        title=clean_title
                    )
                    
                    if hymn_number:
                        self.current_chapter.hymn_number = hymn_number
                    
                    chapter_title_buffer = ""
                    # 继续处理当前行
            
            # 处理读经
            if self.current_chapter and text.startswith('读经：'):
                self.current_chapter.scripture = text.replace('读经：', '').strip()
                continue
            
            # 处理诗歌编号
            if self.current_chapter and ('诗歌：' in text or any(text.startswith(prefix) for prefix in ['JL ', 'MC ', 'SC ', 'RM ', 'NL/', 'HL ', 'RA '])):
                if self.current_chapter.hymn_number:
                    self.current_chapter.hymn_number += ' ' + text
                else:
                    self.current_chapter.hymn_number = text
                continue
            
            # 检测职事信息摘录的开始
            if text == '职事信息摘录：' or text.startswith('职事信息摘录'):
                # 保存之前收集的职事信息（如果有且有效）
                if collecting_ministry and ministry_buffer and self.current_chapter:
                    # 过滤掉无效内容后再保存
                    valid_ministry = [t for t in ministry_buffer if self._is_valid_ministry_text(t)]
                    if valid_ministry:
                        self.current_chapter.ministry_excerpt = '\n\n'.join(valid_ministry)
                # 开始新的职事信息收集
                ministry_buffer = []
                collecting_ministry = True
                # 注意：不再设置current_node = None，让纲目继续解析
                continue
            
            # 如果正在收集职事信息
            if collecting_ministry:
                # 继续收集职事信息内容（跳过空行和无效内容）
                if text and self._is_valid_ministry_text(text):
                    ministry_buffer.append(text)
                continue
            
            # 提取大纲层级(仅标题)
            if self.current_chapter:
                if re.match(r'^[壹贰叁肆伍陆柒捌玖拾]+\s', text):
                    level = self._extract_level_marker(text)
                    title = self._clean_title(text)  # 清理标题，去掉层级标记
                    self.current_level1 = Content(level=level, title=title)
                    self.current_chapter.add_outline_section(self.current_level1)
                    self.current_level2 = None
                    self.current_level3 = None
                    current_node = self.current_level1  # 设置当前节点
                    
                elif re.match(r'^[一二三四五六七八九十]+\s', text) and self.current_level1:
                    level = self._extract_level_marker(text)
                    title = self._clean_title(text)
                    self.current_level2 = Content(level=level, title=title)
                    self.current_level1.add_child(self.current_level2)
                    self.current_level3 = None
                    current_node = self.current_level2  # 设置当前节点
                    
                elif re.match(r'^\d+\s', text) and self.current_level2:
                    level = self._extract_level_marker(text)
                    title = self._clean_title(text)
                    self.current_level3 = Content(level=level, title=title)
                    self.current_level2.add_child(self.current_level3)
                    current_node = self.current_level3  # 设置当前节点
                    
                elif re.match(r'^[a-z]\s', text) and self.current_level3:
                    level = self._extract_level_marker(text)
                    title = self._clean_title(text)
                    level4 = Content(level=level, title=title)
                    self.current_level3.add_child(level4)
                    current_node = level4  # 设置当前节点
            
            # 处理经文内容（verses样式或经文格式）
            if self.current_chapter:
                # 方式1：通过样式名识别 (支持'verses'和'０c 經節')
                is_verse_by_style = para.style and (para.style.name == 'verses' or para.style.name == '０c 經節')
                # 方式2：通过内容格式识别（如：腓2:5	经文内容...）
                is_verse_by_format = self._is_verse_line(text)
                
                if (is_verse_by_style or is_verse_by_format) and text:
                    # 处理"从略"占位符：用缓存的经文范围替换
                    if '从略' in text:
                        verse_info = self._extract_verse_range(text)
                        if verse_info:
                            book, start, end, is_omitted = verse_info
                            if is_omitted:
                                # 从缓存获取该范围的经文
                                cached_verses = self._get_cached_verse_range(book, start, end)
                                if cached_verses:
                                    if self.current_chapter.scripture_verses:
                                        self.current_chapter.scripture_verses += '\n' + cached_verses
                                    else:
                                        self.current_chapter.scripture_verses = cached_verses
                        continue
                    
                    # 正常经文：缓存并添加
                    self._cache_verse(text)
                    verse_text = text.strip()
                    if self.current_chapter.scripture_verses:
                        self.current_chapter.scripture_verses += '\n' + verse_text
                    else:
                        self.current_chapter.scripture_verses = verse_text
        
        # 保存最后收集的职事信息（过滤无效内容）
        if collecting_ministry and ministry_buffer and self.current_chapter:
            # 过滤掉无效内容后再保存
            valid_ministry = [t for t in ministry_buffer if self._is_valid_ministry_text(t)]
            if valid_ministry:
                self.current_chapter.ministry_excerpt = '\n\n'.join(valid_ministry)
        
        if self.current_chapter:
            chapters.append(self.current_chapter)
        
        return chapters
    
    def parse_listen_doc(self, docx_path: str, chapters: List[Chapter]):
        """
        解析听抄文档 - 创建带详细内容的结构
        
        这个文档包含：
        - 与纲目对应的大纲结构（重复）
        - 详细的正文说明
        - 听抄的实际内容
        """
        # print(f"  解析详细内容（{os.path.basename(docx_path)}）...")
        doc = load_document(docx_path)
        current_chapter_num = 0
        self.reset_state()
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # 获取样式类型（支持样式名称映射或直接匹配）
            style_type = self.STYLE_MAP.get(para.style.name) if para.style and para.style.name else None
            
            # 如果样式映射失败，尝试通过文本特征判断
            if not style_type:
                if re.match(r'^第[一二三四五六七八九十]+篇', text):
                    style_type = 'chapter_title'
                elif re.match(r'^[壹贰叁肆伍陆柒捌玖拾]+[、\s]', text):
                    style_type = 'section_level1'
                elif re.match(r'^[一二三四五六七八九十]+[、\s]', text):
                    style_type = 'section_level2'
                elif re.match(r'^\d+[、\s]', text):
                    style_type = 'section_level3'
                else:
                    style_type = 'content'
            
            if style_type == 'chapter_title':
                chapter_num = self._extract_chapter_number(text)
                current_chapter_num = chapter_num
                # 找到对应的章节
                for chapter in chapters:
                    if chapter.number == chapter_num:
                        self.current_chapter = chapter
                        break
                self.current_level1 = None
                self.current_level2 = None
                self.current_level3 = None
                
            elif style_type == 'section_level1' and self.current_chapter:
                # 创建新的大点节点(带内容的)
                level = self._extract_level_marker(text)
                title = self._clean_title(text)
                self.current_level1 = Content(level=level, title=title)
                self.current_chapter.add_detail_section(self.current_level1)
                self.current_level2 = None
                self.current_level3 = None
                    
            elif style_type == 'section_level2' and self.current_level1:
                level = self._extract_level_marker(text)
                title = self._clean_title(text)
                self.current_level2 = Content(level=level, title=title)
                self.current_level1.add_child(self.current_level2)
                self.current_level3 = None
                    
            elif style_type == 'section_level3' and self.current_level2:
                level = self._extract_level_marker(text)
                title = self._clean_title(text)
                self.current_level3 = Content(level=level, title=title)
                self.current_level2.add_child(self.current_level3)
                    
            elif style_type == 'content':
                # 添加正文内容到对应的层级
                if self.current_level3:
                    self.current_level3.add_content(text)
                elif self.current_level2:
                    self.current_level2.add_content(text)
                elif self.current_level1:
                    self.current_level1.add_content(text)
                elif self.current_chapter:
                    # 职事信息内容（在所有大点之外的内容）
                    self.current_chapter.message_content.append(text)
    
    def _detect_level_type(self, text: str) -> str:
        """检测文本的层级类型"""
        if re.match(r'^[壹贰叁肆伍陆柒捌玖拾]+\s', text):
            return 'level1'
        elif re.match(r'^[一二三四五六七八九十]+\s', text):
            return 'level2'
        elif re.match(r'^\d+\s', text):
            return 'level3'
        elif re.match(r'^[a-z]\s', text):
            return 'level4'
        return 'unknown'
    
    def _extract_chapter_number(self, text: str) -> int:
        """从标题中提取篇章编号（支持第1-99篇）"""
        match = re.search(r'第([一二三四五六七八九十百]+|\d+)篇', text)
        if match:
            num_str = match.group(1)
            
            # 如果是阿拉伯数字，直接转换
            if num_str.isdigit():
                return int(num_str)
            
            # 中文数字映射（支持1-99）
            num_map = {
                '一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6, '七': 7, '八': 8, '九': 9, '十': 10,
                '十一': 11, '十二': 12, '十三': 13, '十四': 14, '十五': 15, '十六': 16, '十七': 17, '十八': 18, '十九': 19, '二十': 20,
                '二十一': 21, '二十二': 22, '二十三': 23, '二十四': 24, '二十五': 25, '二十六': 26, '二十七': 27, '二十八': 28, '二十九': 29, '三十': 30
            }
            return num_map.get(num_str, 1)
        return 1
    
    def _extract_level_marker(self, text: str) -> str:
        """提取层级标识"""
        patterns = [
            r'^([壹贰叁肆伍陆柒捌玖拾]+)\s',
            r'^([一二三四五六七八九十]+)\s',
            r'^(\d+)\s',
            r'^([a-z])\s',
        ]
        
        for pattern in patterns:
            match = re.match(pattern, text)
            if match:
                return match.group(1)
        
        return ''
    
    def parse_morning_revival_doc(self, docx_path: str, chapters: List[Chapter]):
        """
        解析晨兴文档 - 按周和按天的晨兴内容
        
        文档结构：
        - 第X周 • 纲目（整周的纲目，不分天）
        - 第X周 • 诗歌（图片）
        - 周一/周二... （每天只有晨兴喂养和信息选读，共享周纲目）
        
        支持.doc和.docx格式，支持基于样式和基于文本的解析
        """
        # print(f"  解析晨兴内容（{os.path.basename(docx_path)}）...")
        
        # 检查文件格式
        is_doc_format = docx_path.endswith('.doc') and not docx_path.endswith('.docx')
        
        # 加载文档进行文本解析
        doc = load_document(docx_path)
        
        # 对于.doc文件，需要先转换为临时.docx才能提取图片
        # 使用LibreOffice转换（跨平台方案）
        temp_docx_path = None
        if is_doc_format:
            import tempfile
            import subprocess
            import shutil
            
            temp_dir = tempfile.mkdtemp()
            abs_path = os.path.abspath(docx_path)
            
            try:
                # 尝试找到LibreOffice/soffice命令
                soffice_commands = [
                    'soffice',
                    'libreoffice',
                    r'C:\Program Files\LibreOffice\program\soffice.exe',
                    r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
                ]
                
                soffice_path = None
                for cmd in soffice_commands:
                    if shutil.which(cmd) or os.path.exists(cmd):
                        soffice_path = cmd
                        break
                
                if soffice_path:
                    # 使用LibreOffice转换
                    result = subprocess.run(
                        [soffice_path, '--headless', '--convert-to', 'docx', '--outdir', temp_dir, abs_path],
                        capture_output=True,
                        timeout=30
                    )
                    
                    if result.returncode == 0:
                        docx_name = os.path.splitext(os.path.basename(docx_path))[0] + '.docx'
                        temp_docx_path = os.path.join(temp_dir, docx_name)
                        
                        if os.path.exists(temp_docx_path):
                            # 从临时docx提取图片
                            print("  提取诗歌图片...")
                            doc_id = os.path.basename(docx_path).replace('.doc', '').replace('.docx', '')
                            self._extract_hymn_images(temp_docx_path, chapters, doc_id)
                            temp_docx_path = temp_docx_path  # 保存路径供后续使用
                        else:
                            print("    ⚠ LibreOffice转换失败，跳过图片提取")
                else:
                    print("    ⚠ 未找到LibreOffice，跳过图片提取")
                    print("    提示: 安装LibreOffice或手动将.doc转换为.docx")
                
            except subprocess.TimeoutExpired:
                print("    ⚠ LibreOffice转换超时")
            except Exception as e:
                print(f"    ⚠ 提取诗歌图片失败: {e}")
            finally:
                # 清理临时文件
                shutil.rmtree(temp_dir, ignore_errors=True)
        else:
            # .docx文件可以直接提取图片
            print("  提取诗歌图片...")
            doc_id = os.path.basename(docx_path).replace('.doc', '').replace('.docx', '')
            self._extract_hymn_images(docx_path, chapters, doc_id)
        
        # 统计样式使用情况，判断使用哪种解析策略
        style_counts = {}
        for para in doc.paragraphs[:200]:  # 检查前200段
            if para.text.strip():
                if para.style and para.style.name:
                    style_counts[para.style.name] = style_counts.get(para.style.name, 0) + 1
        
        # 夏季样式标记
        has_summer_styles = any(s in style_counts for s in ['第一周', '第一周右', '周期', '１綱要大點壹'])
        
        if has_summer_styles:
            # 使用基于样式的解析（夏季）
            self._parse_morning_revival_by_styles(doc, chapters)
        else:
            # 使用基于文本的解析（秋季）
            paragraphs = [para.text.strip() for para in doc.paragraphs]
            self._parse_morning_revival_by_text(paragraphs, chapters)
    
    def _parse_morning_revival_by_text(self, paragraphs: List[str], chapters: List[Chapter]):
        """
        基于文本模式的晨兴文档解析（秋季 .docx）
        """
        
        current_week = 0
        current_day = None
        current_day_num = 0
        current_section = None
        content_buffer = []
        current_revival = None
        day_outlines = {}
        current_day_outline = []
        hymn_buffer = []
        seen_revival_days = set()
        
        in_content = False  # 是否已进入正文内容
        in_outline_section = False  # 是否在纲目部分
        
        line_num = 0
        debug_log = []
        debug_log.append(f"=== _parse_morning_revival_by_text started with {len(paragraphs)} paragraphs ===")
        
        for text in paragraphs:
            line_num += 1
            if not text:
                continue
            
            # 跳过目录
            if not in_content:
                if text.startswith('第一周') and '纲目' in text:
                    in_content = True
                    # 同时设置in_outline_section,因为周一周二的纲目紧跟在后面
                    in_outline_section = True
                    # 设置current_week=1,避免第二次遇到"第一周 • 纲目"时触发周切换
                    current_week = 1
                    current_section = 'week_outline'
                    debug_log.append(f"Line {line_num}: First week outline marker detected, starting week 1")
                continue
            
            # 检测新的周（第X周 • 纲目）
            week_match = re.match(r'第([一二三四五六七八九])周\s*[•·]\s*纲目', text)
            if week_match:
                week_map = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6, '七': 7, '八': 8, '九': 9}
                new_week = week_map.get(week_match.group(1), 0)
                
                # 只有当周号变化时才处理（避免重复的页眉）
                if new_week != current_week:
                    debug_log.append(f"Line {line_num}: Detected week change from {current_week} to {new_week}, day_outlines keys: {list(day_outlines.keys())}")
                    # 保存上一周的数据
                    if current_week > 0 and current_week <= len(chapters):
                        # 保存最后一天的晨兴内容
                        if current_revival:
                            if content_buffer:
                                if current_section == 'feeding':
                                    current_revival.morning_feeding = content_buffer.copy()
                                elif current_section == 'reading':
                                    current_revival.message_reading = content_buffer.copy()
                            chapters[current_week - 1].morning_revivals.append(current_revival)
                            debug_log.append(f"Line {line_num}: Saved last day revival at week switch")
                            current_revival = None
                            content_buffer = []
                        
                        # 保存最后一天的纲目
                        if current_day_outline and current_day_num > 0:
                            day_outlines[current_day_num] = current_day_outline.copy()
                            debug_log.append(f"Line {line_num}: Saved last day {current_day_num} outline at week switch")
                        
                        # 保存诗歌信息
                        if hymn_buffer:
                            chapters[current_week - 1].hymn_number = '\n'.join(hymn_buffer)
                            hymn_buffer = []
                        
                        # 将每天的纲目保存到临时位置（稍后与晨兴内容合并）
                        if day_outlines:
                            chapters[current_week - 1]._day_outlines = day_outlines.copy()
                            debug_log.append(f"Line {line_num}: Saved week {current_week} day_outlines: {list(day_outlines.keys())}")
                            day_outlines = {}
                    
                    current_week = new_week
                    current_section = 'week_outline'
                    in_outline_section = True
                    current_day_num = 0
                    current_day_outline = []
                    debug_log.append(f"Line {line_num}: Starting new week {current_week}")
                else:
                    # 遇到重复的周纲目标记(页眉),不做任何处理,继续收集纲目
                    # 注意:不要清空current_day_outline,因为后续内容属于同一天
                    debug_log.append(f"Line {line_num}: Skip duplicate week outline marker for week {current_week}, continuing day {current_day_num}")
                continue
            
            # 检测诗歌部分（第X周 • 诗歌）
            hymn_match = re.match(r'第([一二三四五六七八九])周\s*[•·]\s*诗歌', text)
            if hymn_match:
                in_outline_section = False
                current_section = 'hymn'
                # 保存最后一天的纲目
                if current_day_outline and current_day_num > 0:
                    day_outlines[current_day_num] = current_day_outline.copy()
                    debug_log.append(f"Line {line_num}: Saved outline for day {current_day_num} at hymn section, {len(current_day_outline)} lines")
                    current_day_outline = []
                debug_log.append(f"Line {line_num}: At hymn section, day_outlines keys: {list(day_outlines.keys())}")
                continue
            
            # 在纲目部分检测天标记(周一、周二等)
            # 注意: 使用[\s\u3000]*匹配普通空格和全角空格
            if in_outline_section:
                day_in_outline_match = re.match(r'周[\s\u3000]*([一二三四五六])', text)
                if day_in_outline_match:
                    # 保存前一天的纲目
                    if current_day_outline and current_day_num > 0:
                        day_outlines[current_day_num] = current_day_outline.copy()
                        debug_log.append(f"Line {line_num}: Saved outline for day {current_day_num}, {len(current_day_outline)} lines, day_outlines keys now: {list(day_outlines.keys())}")
                    elif current_day_num > 0:
                        debug_log.append(f"Line {line_num}: NOT saving day {current_day_num} - outline is empty, current_day_outline length: {len(current_day_outline)}")
                    
                    # 开始新的一天
                    day_map = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6}
                    current_day_num = day_map.get(day_in_outline_match.group(1), 0)
                    current_day_outline = []
                    debug_log.append(f"Line {line_num}: Start collecting outline for day {current_day_num}, text=[{text}]")
                    continue
            
            # 检测申言部分（第X周 • 申言）- 跳过不处理
            proclamation_match = re.match(r'第([一二三四五六七八九])周\s*[•·]\s*申言', text)
            if proclamation_match:
                in_outline_section = False
                current_section = 'proclamation'  # 标记为申言部分，后续内容将被跳过
                continue
            
            # 检测晨兴内容部分的天（必须是"第X周 • 周X"格式）
            # 注意：这个标记既标志着纲目部分的结束，也标志着晨兴内容的开始
            day_match = re.match(r'第[一二三四五六七八九]周\s*[•·]\s*周\s*([一二三四五六七])', text)
            if day_match and current_week > 0:
                debug_log.append(f"Line {line_num}: [Text: {text}], [in_outline: {in_outline_section}]")
                # 检查是否重复（避免页眉重复创建）
                day_key = f"{current_week}_{text}"
                if day_key in seen_revival_days:
                    debug_log.append(f"Line {line_num}: Skip duplicate day [{text}]")
                    continue
                
                seen_revival_days.add(day_key)
                # 保存当前正在收集的纲目（如果还在纲目部分）
                if in_outline_section and current_day_outline and current_day_num > 0:
                    day_outlines[current_day_num] = current_day_outline.copy()
                    debug_log.append(f"Line {line_num}: Saved outline for day {current_day_num} before exiting outline section, {len(current_day_outline)} lines")
                    current_day_outline = []
                
                # 进入晨兴内容部分，退出纲目部分（无论之前in_outline_section是什么）
                in_outline_section = False
                debug_log.append(f"Line {line_num}: Found revival day [{text}], exiting outline section")
                # 保存前一天的内容
                if current_revival:
                    if content_buffer:
                        if current_section == 'feeding':
                            current_revival.morning_feeding = content_buffer.copy()
                        elif current_section == 'reading':
                            current_revival.message_reading = content_buffer.copy()
                    # 添加到对应章节
                    if current_week <= len(chapters):
                        chapters[current_week - 1].morning_revivals.append(current_revival)
                
                # 开始新的一天（每天使用周纲目，不单独解析纲目）
                current_day = text
                current_revival = MorningRevival(day=current_day)
                current_section = None  # 等待晨兴喂养标记
                content_buffer = []
                continue
            
            # 检测晨兴喂养
            if text == '晨兴喂养' or text.startswith('晨兴喂养'):
                current_section = 'feeding'
                content_buffer = []
                continue
            
            # 检测信息选读
            if text == '信息选读' or text.startswith('信息选读'):
                # 保存前面的晨兴喂养内容
                if current_revival and content_buffer:
                    current_revival.morning_feeding = content_buffer.copy()
                current_section = 'reading'
                content_buffer = []
                continue
            
            # 收集内容
            if current_section == 'proclamation':
                # 跳过申言部分的内容
                continue
            elif current_section == 'hymn':
                # 收集诗歌内容
                hymn_buffer.append(text)
            elif in_outline_section and current_day_num > 0:
                # 收集当天的纲目内容
                current_day_outline.append(text)
                if len(current_day_outline) <= 3:  # 只记录前3行避免日志过多
                    debug_log.append(f"Line {line_num}: Collected outline for day {current_day_num}: {text[:50]}")
            elif current_revival and current_section in ['feeding', 'reading']:
                # 处理跨页连接：如果当前内容和前一行内容需要连接
                if content_buffer and self._should_merge_with_previous(content_buffer[-1], text):
                    # 将当前文本连接到前一行，不插入换行
                    content_buffer[-1] = content_buffer[-1] + text
                else:
                    content_buffer.append(text)
        
        # 保存最后一天的内容
        if current_revival:
            if content_buffer:
                if current_section == 'feeding':
                    current_revival.morning_feeding = content_buffer.copy()
                elif current_section == 'reading':
                    current_revival.message_reading = content_buffer.copy()
            if current_week > 0 and current_week <= len(chapters):
                chapters[current_week - 1].morning_revivals.append(current_revival)
        
        # 保存最后一周的数据
        if current_week > 0 and current_week <= len(chapters):
            # 保存最后一天的纲目
            if current_day_outline and current_day_num > 0:
                day_outlines[current_day_num] = current_day_outline.copy()
            
            # 保存诗歌信息
            if hymn_buffer:
                chapters[current_week - 1].hymn_number = '\n'.join(hymn_buffer)
            
            # 保存每天的纲目
            if day_outlines:
                chapters[current_week - 1]._day_outlines = day_outlines.copy()
        
        # 将每天的纲目分配到对应的MorningRevival对象
        for chapter_idx, chapter in enumerate(chapters, 1):
            if hasattr(chapter, '_day_outlines') and chapter._day_outlines:
                debug_log.append(f"Chapter {chapter_idx} has _day_outlines: {list(chapter._day_outlines.keys())}, morning_revivals count: {len(chapter.morning_revivals)}")
                
                # 如果缺失开头的day,用第一个可用的纲目填充
                # (处理第一周的周一周二在文档中缺失纲目标记的情况)
                if chapter._day_outlines:
                    first_available_day = min(chapter._day_outlines.keys())
                    first_outline = chapter._day_outlines[first_available_day]
                    for day in range(1, first_available_day):
                        if day not in chapter._day_outlines:
                            chapter._day_outlines[day] = first_outline
                            debug_log.append(f"Backfilled outline for Chapter {chapter_idx}, Day {day} using Day {first_available_day}'s outline")
                
                # 为每一天分配纲目,如果某天没有纲目则使用前一天的
                last_outline_lines = None
                for day_num, revival in enumerate(chapter.morning_revivals, 1):
                    if day_num in chapter._day_outlines:
                        outline_lines = chapter._day_outlines[day_num]
                        last_outline_lines = outline_lines  # 记录最后一次的纲目
                        debug_log.append(f"Assigning outline to Chapter {chapter_idx}, Day {day_num}, {len(outline_lines)} lines")
                        # 特别针对第一章第六天添加详细debug
                        if chapter_idx == 1 and day_num == 6:
                            debug_log.append(f"=== Special debug for Chapter 1, Day 6 ===")
                            debug_log.append(f"Outline lines: {outline_lines}")
                            for i, line in enumerate(outline_lines):
                                debug_log.append(f"  Line {i}: [{line}]")
                        revival.outline = self._parse_outline_content(outline_lines)
                    elif last_outline_lines:
                        # 使用前一天的纲目
                        debug_log.append(f"Using previous outline for Chapter {chapter_idx}, Day {day_num}, {len(last_outline_lines)} lines")
                        revival.outline = self._parse_outline_content(last_outline_lines)
                    else:
                        debug_log.append(f"No outline for Chapter {chapter_idx}, Day {day_num}")
            else:
                debug_log.append(f"Chapter {chapter_idx} has no _day_outlines or empty, morning_revivals count: {len(chapter.morning_revivals)}")
        
        # 写入调试日志
        debug_log.append("=== Debug log end ===")  # 强制添加一些内容
        if debug_log:
            with open('debug_morning_revival.log', 'w', encoding='utf-8') as f:
                f.write('\n'.join(debug_log))
    
    def _parse_morning_revival_by_styles(self, doc, chapters: List[Chapter]):
        """
        基于样式的晨兴文档解析（夏季 .doc）
        
        文档结构：
        1. 第X周 • 纲目 (周纲目标记)
        2. 周期: 周一 (标记接下来纲目属于周一)
        3. 纲目内容 (多行纲目样式)
        4. 周期: 周二、周三 (标记接下来纲目属于周二、三)
        5. 纲目内容
        6. 第X周 • 周一 (天标记，创建MorningRevival)
        7. 晨兴喂养 (喂养内容)
        8. 信息选读 (选读内容)
        9. 第X周 • 周二 (下一天)
        ...
        
        关键：纲目在"周期"和"第X周•周X"之间，需要先收集再分配
        """
        current_week = 0
        week_to_chapter_map = {}  # 周数 -> 章节索引的映射
        chapter_index = 0  # 当前章节索引（0-based）
        
        # 检查是否有已存在的晨兴内容（即这是第二个晨兴文档）
        existing_chapters_with_revivals = sum(1 for c in chapters if c.morning_revivals)
        if existing_chapters_with_revivals > 0:
            chapter_index = existing_chapters_with_revivals
        
        original_week_to_new_week = {}  # 原始周数到新周数的映射（夏季训练重新编号）
        current_section = None
        current_revival = None
        all_weeks_outlines = {}  # 保存所有周的纲目: {week_num: {day_key: outline_lines}}
        day_outlines = {}  # 天标记 -> 纲目列表的映射 (如 "周一" -> [纲目])
        current_day_key = None  # 当前"周期"对应的天(如"周一"或"周二、周三")
        pending_outline = []  # 当前收集的纲目
        content_buffer = []
        week_hymns = []  # 诗歌信息缓冲
        week_ministry = []  # 职事信息缓冲
        in_outline_section = False  # 是否在纲目收集状态
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            style = para.style.name
            
            # 检测周纲目（第X周 • 纲目）
            # 只处理"第一周"样式，忽略"第一周右"
            if (style == '第一周' and '纲目' in text) or re.match(r'第.+周.*纲目', text):
                # 提取周数（支持中文数字和阿拉伯数字）
                week_match = re.search(r'第([一二三四五六七八九十百]+|\d+)周', text)
                if week_match:
                    week_str = week_match.group(1)
                    
                    # 如果是阿拉伯数字，直接转换
                    if week_str.isdigit():
                        new_week = int(week_str)
                    else:
                        # 中文数字映射（扩展到50周）
                        week_map = {
                            '一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6, '七': 7, '八': 8, '九': 9, '十': 10,
                            '十一': 11, '十二': 12, '十三': 13, '十四': 14, '十五': 15, '十六': 16, '十七': 17, '十八': 18, '十九': 19, '二十': 20,
                            '二十一': 21, '二十二': 22, '二十三': 23, '二十四': 24, '二十五': 25, '二十六': 26, '二十七': 27, '二十八': 28, '二十九': 29, '三十': 30,
                            '三十一': 31, '三十二': 32, '三十三': 33, '三十四': 34, '三十五': 35, '三十六': 36, '三十七': 37, '三十八': 38, '三十九': 39, '四十': 40,
                            '四十一': 41, '四十二': 42, '四十三': 43, '四十四': 44, '四十五': 45, '四十六': 46, '四十七': 47, '四十八': 48, '四十九': 49, '五十': 50
                        }
                        new_week = week_map.get(week_str, 0)
                    
                    # 【夏季训练重新编号】如果是夏季训练（周数较大），重新从1开始编号
                    if new_week >= 20:  # 假设周数>=20的是夏季训练
                        if new_week not in original_week_to_new_week:
                            # 如果是第二个文档，从已有章节数+1开始编号
                            renumbered_week = len(original_week_to_new_week) + 1 + chapter_index
                            original_week_to_new_week[new_week] = renumbered_week
                        new_week = original_week_to_new_week[new_week]

                    if new_week > 0 and new_week != current_week:
                        # 【重要修复】保存上一周的纲目到全局映射中（累积保存，不覆盖）
                        if current_week in week_to_chapter_map and day_outlines:
                            if current_week in all_weeks_outlines:
                                # 合并已有纲目，不覆盖
                                all_weeks_outlines[current_week].update(day_outlines)
                            else:
                                # 新建纲目映射
                                all_weeks_outlines[current_week] = day_outlines.copy()
                        
                        # 保存上一周最后一组纲目（非常重要！在清空day_outlines之前保存）
                        if current_day_key and pending_outline:
                            clean_key = current_day_key.replace('\u3000', '').replace(' ', '')
                            days = re.findall(r'周[一二三四五六]', clean_key)
                            for day in days:
                                day_outlines[day] = pending_outline.copy()
                        
                        # 【重要修复】再次保存（包含最后一组纲目）（累积保存）
                        if current_week in week_to_chapter_map and day_outlines:
                            if current_week in all_weeks_outlines:
                                # 合并已有纲目，不覆盖
                                all_weeks_outlines[current_week].update(day_outlines)
                            else:
                                # 新建纲目映射
                                all_weeks_outlines[current_week] = day_outlines.copy()
                        
                        # 保存上一周的数据
                        if current_revival and current_week in week_to_chapter_map:
                            if content_buffer:
                                if current_section == 'feeding':
                                    current_revival.morning_feeding = content_buffer
                                elif current_section == 'reading':
                                    current_revival.message_reading = content_buffer
                            chap_idx = week_to_chapter_map[current_week]
                            if chap_idx < len(chapters):
                                chapters[chap_idx].morning_revivals.append(current_revival)
                        
                        # 保存上一周的诗歌和职事信息
                        if current_week in week_to_chapter_map:
                            chap_idx = week_to_chapter_map[current_week]
                            if chap_idx < len(chapters):
                                if week_hymns:
                                    chapters[chap_idx].hymn_number = '\n'.join(week_hymns)
                                if week_ministry:
                                    # 过滤无效内容后再保存
                                    valid_ministry = [t for t in week_ministry if self._is_valid_ministry_text(t)]
                                    if valid_ministry:
                                        chapters[chap_idx].ministry_excerpt = '\n\n'.join(valid_ministry)
                        
                        # 开始新周，映射到下一个chapter
                        current_week = new_week
                        if current_week not in week_to_chapter_map:
                            week_to_chapter_map[current_week] = chapter_index
                            chapter_index += 1
                        
                        # 清空上一周的纲目映射，准备收集新周的纲目
                        day_outlines = {}
                        current_day_key = None
                        pending_outline = []
                        
                        week_hymns = []  # 诗歌信息缓冲
                        week_ministry = []  # 职事信息缓冲
                        current_revival = None
                        current_day_outline = []
                        in_outline_section = True  # 进入纲目部分
                        current_section = 'outline'
                continue
            
            # 检测"周期"样式 - 标记接下来的纲目属于哪些天
            if style == '周期':
                # 保存上一组纲目
                if current_day_key and pending_outline:
                    # 如果是"周二、周三"这种,需要分别存储
                    # 先去除全角空格,统一格式
                    clean_key = current_day_key.replace('\u3000', '').replace(' ', '')
                    days = re.findall(r'周[一二三四五六]', clean_key)
                    for day in days:
                        day_outlines[day] = pending_outline.copy()
                
                # 记录新的天标记,开始收集新纲目
                current_day_key = text  # 如"周一"或"周二、周三"
                pending_outline = []
                in_outline_section = True
                current_section = 'outline'
                continue
            
            # 检测具体的天标记（第X周 • 周一/周二等）
            # 只处理"第一周"样式，忽略"第一周右"（右列是同一天的延续）
            if style == '第一周' and re.search(r'周[一二三四五六七]', text) and '纲目' not in text and '诗歌' not in text:
                # 提取周数进行重新编号
                week_match = re.search(r'第([一二三四五六七八九十百]+|\d+)周', text)
                if week_match:
                    week_str = week_match.group(1)
                    original_week = int(week_str) if week_str.isdigit() else {
                        '一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6, '七': 7, '八': 8, '九': 9, '十': 10,
                        '十一': 11, '十二': 12, '十三': 13, '十四': 14, '十五': 15, '十六': 16, '十七': 17, '十八': 18, '十九': 19, '二十': 20,
                        '二十一': 21, '二十二': 22, '二十三': 23, '二十四': 24, '二十五': 25, '二十六': 26, '二十七': 27, '二十八': 28, '二十九': 29, '三十': 30,
                    }.get(week_str, 0)
                    
                    # 夏季训练重新编号
                    if original_week >= 20:
                        if original_week not in original_week_to_new_week:
                            # 如果是第二个文档，从已有章节数+1开始编号
                            renumbered_week = len(original_week_to_new_week) + 1 + chapter_index
                            original_week_to_new_week[original_week] = renumbered_week
                        mapped_week = original_week_to_new_week[original_week]
                    else:
                        mapped_week = original_week
                # 保存上一天的数据
                if current_revival and current_week in week_to_chapter_map:
                    # 保存上一天的喂养/选读
                    if content_buffer:
                        if current_section == 'feeding':
                            current_revival.morning_feeding = content_buffer
                        elif current_section == 'reading':
                            current_revival.message_reading = content_buffer
                    chap_idx = week_to_chapter_map[current_week]
                    if chap_idx < len(chapters):
                        chapters[chap_idx].morning_revivals.append(current_revival)
                
                # 【重要修复】保存上一组纲目（如果还在收集中，例如周六纲目）
                if current_day_key and pending_outline:
                    clean_key = current_day_key.replace('\u3000', '').replace(' ', '')
                    days = re.findall(r'周[一二三四五六]', clean_key)
                    for day in days:
                        day_outlines[day] = pending_outline.copy()
                    pending_outline = [] # 清空已保存的纲目

                # 创建新的一天,从day_outlines获取对应纲目
                current_revival = MorningRevival(day=text)
                # 提取天标记中的"周X"部分
                day_match = re.search(r'周[一二三四五六]', text)
                if day_match:
                    day_key = day_match.group()
                    if day_key in day_outlines:
                        current_revival.outline = self._parse_outline_content(day_outlines[day_key])
                
                # 退出纲目收集状态，准备收集喂养内容
                in_outline_section = False
                current_section = None
                content_buffer = []
                continue
            
            # 忽略"第一周右"样式 - 这是分页延续，不是新的天标记
            if style == '第一周右' and re.search(r'周[一二三四五六七]', text):
                # 这是跨页的天标记重复，跳过不处理
                continue
            
            # 检测喂养样式标记（切换到喂养内容收集）
            if style == '喂养选读' and '晨兴喂养' in text:
                in_outline_section = False
                current_section = 'feeding'
                content_buffer = []
                continue
            
            # 检测参读/选读
            if style == '喂养选读' and ('信息选读' in text or '选读' in text):
                if current_revival and content_buffer and current_section == 'feeding':
                    current_revival.morning_feeding = content_buffer
                current_section = 'reading'
                content_buffer = []
                continue
            
            if style == '参读光亮':
                if current_revival and content_buffer:
                    if current_section == 'feeding':
                        current_revival.morning_feeding = content_buffer
                    elif current_section == 'reading':
                        current_revival.message_reading = content_buffer
                current_section = None
                content_buffer = []
                continue
            
            # 收集内容
            if in_outline_section and current_section == 'outline':
                # 收集纲目内容（在"周期"和"第X周•周X"之间）
                if style in ['１綱要大點壹', '２綱要中點一', '３綱要小點１', '４綱要分點ａ'] or \
                   re.match(r'^[壹贰叁肆伍陆柒捌玖拾一二三四五六七八九十\d]+[、\s]', text):
                    pending_outline.append(text)
                # 检测诗歌（在纲目部分）
                elif 'EM' in text or 'RK' in text or '诗歌' in text:
                    week_hymns.append(text)
            elif current_section in ['feeding', 'reading']:
                # 收集喂养/选读内容，但要过滤掉周标记文本（如"第二十五周 • 周一"）
                # 这些文本是分页时的标记，不是实际内容
                if not re.match(r'^第.+周\s*[•·]\s*周[一二三四五六七]$', text):
                    # 处理跨页连接：如果当前内容和前一行内容需要连接
                    if content_buffer and self._should_merge_with_previous(content_buffer[-1], text):
                        # 将当前文本连接到前一行，不插入换行
                        content_buffer[-1] = content_buffer[-1] + text
                    else:
                        content_buffer.append(text)
            elif current_section is None and current_week in week_to_chapter_map:
                # 职事信息（在纲目之后、喂养之前的内容）
                if style in ['信息正文18', '信息正文17', '信息正文16', '信息正文3'] and len(text) > 20:
                    week_ministry.append(text)
        
        # 保存最后一组纲目(如果还在收集中)
        if current_day_key and pending_outline:
            clean_key = current_day_key.replace('\u3000', '').replace(' ', '')
            days = re.findall(r'周[一二三四五六七]', clean_key)
            for day in days:
                day_outlines[day] = pending_outline.copy()
        
        # 保存最后一天的内容
        if current_revival and current_week in week_to_chapter_map:
            # 保存最后一天的喂养/选读
            if content_buffer:
                if current_section == 'feeding':
                    current_revival.morning_feeding = content_buffer
                elif current_section == 'reading':
                    current_revival.message_reading = content_buffer
            chap_idx = week_to_chapter_map[current_week]
            if chap_idx < len(chapters):
                chapters[chap_idx].morning_revivals.append(current_revival)
                
                # 保存最后一周的诗歌和职事信息
                if week_hymns:
                    chapters[chap_idx].hymn_number = '\n'.join(week_hymns)
                if week_ministry:
                    # 过滤无效内容后再保存
                    valid_ministry = [t for t in week_ministry if self._is_valid_ministry_text(t)]
                    if valid_ministry:
                        chapters[chap_idx].ministry_excerpt = '\n\n'.join(valid_ministry)
        
        # 保存最后一周的纲目
        if current_day_key and pending_outline:
            clean_key = current_day_key.replace('\u3000', '').replace(' ', '')
            days = re.findall(r'周[一二三四五六七]', clean_key)
            for day in days:
                day_outlines[day] = pending_outline.copy()
        
        # 【重要修复】保存最后一周的纲目到全局映射中（累积保存）
        if current_week in week_to_chapter_map and day_outlines:
            if current_week in all_weeks_outlines:
                # 合并已有纲目，不覆盖
                all_weeks_outlines[current_week].update(day_outlines)
            else:
                # 新建纲目映射
                all_weeks_outlines[current_week] = day_outlines.copy()
        
        # 【重要修复】为所有章节分配纲目，不仅仅是最后一章
        for chapter_idx, chapter in enumerate(chapters, 1):
            
            # 显示每个revival的基本信息
            for day_num, revival in enumerate(chapter.morning_revivals, 1):
                day_text = revival.day.replace('•', ' ')  # 避免编码问题
                # print(f"    Day {day_num}: {day_text} -> has outline: {bool(revival.outline)}")
                if revival.outline:
                    pass  # print(f"      Outline content count: {len(revival.outline)}")
                else:
                    pass  # print(f"      No outline content")
            
            # 检查该章节对应的周是否有纲目
            # 对于夏季训练，需要使用week_to_chapter_map的反向映射找到正确的周数
            week_num = None
            for week, chap_idx in week_to_chapter_map.items():
                if chap_idx == chapter_idx - 1:  # 修复：chap_idx是0-based，chapter_idx是1-based
                    week_num = week
                    break
            
            if week_num is None:
                week_num = chapter_idx  # 回退策略：第X章对应第X周
                
            if week_num in all_weeks_outlines:
                week_day_outlines = all_weeks_outlines[week_num]
                
                # 将day_outlines保存到chapter对象
                converted_outlines = {}
                day_map = {'周一': 1, '周二': 2, '周三': 3, '周四': 4, '周五': 5, '周六': 6, '周日': 7}
                for day_key, outline_lines in week_day_outlines.items():
                    if day_key in day_map:
                        converted_outlines[day_map[day_key]] = outline_lines
                        # print(f"    Converted {day_key} -> day {day_map[day_key]}: {len(outline_lines)} lines")
                
                if converted_outlines:
                    chapter._day_outlines = converted_outlines
                    
                    # 为每一天分配纲目，如果某天没有纲目则使用前一天的
                    last_outline_lines = None
                    for day_num, revival in enumerate(chapter.morning_revivals, 1):
                        if day_num in converted_outlines:
                            outline_lines = converted_outlines[day_num]
                            last_outline_lines = outline_lines
                            parsed_outline = self._parse_outline_content(outline_lines)
                            revival.outline = parsed_outline
                            
                            # 特别针对第一章第六天添加详细debug
                            if chapter_idx == 1 and day_num == 6:
                                pass  # for i, content in enumerate(parsed_outline):
                                    # print(f"      Content {i}: {content.level} - {content.title[:50]}")
                        elif last_outline_lines:
                            revival.outline = self._parse_outline_content(last_outline_lines)
            
            # 特别处理第一章
            if chapter_idx == 1:
                for day_num, revival in enumerate(chapter.morning_revivals, 1):
                    outline_count = len(revival.outline) if revival.outline else 0
                    # print(f"    Day {day_num} ({revival.day.replace('•', ' ')}): {outline_count} outline items")
    
    def _parse_outline_content(self, content_lines: List[str]) -> List[Content]:
        """
        解析晨兴中的纲目内容，转换为Content对象列表
        
        Args:
            content_lines: 纲目文本行列表
            
        Returns:
            Content对象列表
        """
        # 添加调试信息
        debug_msg = f"_parse_outline_content called with {len(content_lines)} lines"
        if content_lines:
            debug_msg += f", first line: {content_lines[0][:50]}..."
        with open('debug_outline_parsing.log', 'a', encoding='utf-8') as f:
            f.write(debug_msg + '\n')
            for i, line in enumerate(content_lines):
                f.write(f"  Line {i}: {line}\n")
            
            # 特别检查是否包含周六特征内容
            is_saturday_content = any(keyword in ' '.join(content_lines) for keyword in ["贰", "四", "五"])
            saturday_keywords = ["四　我们必须在新样里", "五　我们日日得更新", "六　我们调和的灵需要扩展"]
            for keyword in saturday_keywords:
                for line in content_lines:
                    if keyword in line:
                        f.write(f"*** FOUND SATURDAY KEYWORD: {keyword} in line: {line}\n")
                        break
        
        result = []
        current_level1 = None
        current_level2 = None
        current_level3 = None
        current_level4 = None
        
        for text in content_lines:
            if not text or len(text) < 2:
                continue
            
            # 检测各级纲目 - 使用更灵活的匹配方式
            # 壹、贰、叁等 (level1)
            level1_match = re.match(r'^([壹贰叁肆伍陆柒捌玖拾一二三四五六七八九十])[\s　]+(.*)', text)
            if level1_match:
                level = level1_match.group(1)
                title = level1_match.group(2)
                current_level1 = Content(level=level, title=title)
                result.append(current_level1)
                current_level2 = None
                current_level3 = None
                current_level4 = None
                continue
            
            # 一、二、三等 (level2)
            level2_match = re.match(r'^([一二三四五六七八九十百]+)[\s　]+(.*)', text)
            if level2_match:
                level = level2_match.group(1)
                title = level2_match.group(2)
                current_level2 = Content(level=level, title=title)
                if current_level1:
                    current_level1.add_child(current_level2)
                else:
                    # 如果没有level1父级,直接作为顶级添加
                    result.append(current_level2)
                current_level3 = None
                current_level4 = None
                continue
            
            # 1、2、3等 (level3)
            level3_match = re.match(r'^(\d+)[\s　]+(.*)', text)
            if level3_match:
                level = level3_match.group(1)
                title = level3_match.group(2)
                current_level3 = Content(level=level, title=title)
                if current_level2:
                    current_level2.add_child(current_level3)
                elif current_level1:
                    current_level1.add_child(current_level3)
                else:
                    # 如果没有父级,直接作为顶级添加
                    result.append(current_level3)
                current_level4 = None
                continue
            
            # a、b、c等 (level4)
            level4_match = re.match(r'^([a-z])[\s　]+(.*)', text)
            if level4_match and current_level3:
                level = level4_match.group(1)
                title = level4_match.group(2)
                current_level4 = Content(level=level, title=title)
                current_level3.add_child(current_level4)
                continue
            
            # 如果是段落文本，添加到最近的节点
            if current_level4:
                current_level4.add_content(text)
            elif current_level3:
                current_level3.add_content(text)
            elif current_level2:
                current_level2.add_content(text)
            elif current_level1:
                current_level1.add_content(text)
        
        return result
    

    def _read_doc_with_win32com(self, doc_path: str, return_doc: bool = False):
        """
        使用win32com读取.doc格式文档
        
        Args:
            doc_path: .doc文件路径
            return_doc: 是否返回文档对象（用于提取图片）
            
        Returns:
            如果return_doc=True: (段落文本列表, 文档对象)
            否则: 段落文本列表
        """
        import os
        try:
            import win32com.client
        except ImportError:
            print("    ⚠ 错误：需要安装pywin32库来读取.doc格式")
            print("    请运行: pip install pywin32")
            return [] if not return_doc else ([], None)
        
        # 转换为绝对路径
        abs_path = os.path.abspath(doc_path)
        
        paragraphs = []
        word = None
        doc = None
        
        try:
            # 创建Word应用程序对象
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            # 打开文档
            doc = word.Documents.Open(abs_path)
            
            # 读取所有段落
            for para in doc.Paragraphs:
                text = para.Range.Text.strip()
                if text:
                    paragraphs.append(text)
            
            if return_doc:
                return paragraphs, doc  # 返回文档对象用于后续操作
            
        except Exception as e:
            # print(f"    ✗ 读取.doc文件失败: {e}")
            return [] if not return_doc else ([], None)
        finally:
            # 如果不需要返回文档对象，则关闭
            if not return_doc:
                if doc:
                    doc.Close(False)
                if word:
                    word.Quit()
        
        return paragraphs
    
    def _extract_hymn_images(self, doc_or_docx, chapters: List[Chapter], doc_identifier: str = ""):
        """
        从Word文档中提取诗歌图片并保存（跨平台方法）
        
        Args:
            doc_or_docx: docx文件路径
            chapters: 章节列表
            doc_identifier: 文档标识符（用于区分多文档时的图片来源）
        """
        import os
        import zipfile
        from io import BytesIO
        
        try:
            from PIL import Image
            from docx import Document as DocxDocument
            
            # 创建图片输出目录
            output_dir = os.path.join(self.output_dir, 'images')
            os.makedirs(output_dir, exist_ok=True)
            
            docx_path = doc_or_docx
            if not os.path.exists(docx_path):
                print(f"    ⚠ 文件不存在: {docx_path}")
                return
            
            # 先读取文档,识别其中包含的周数
            try:
                doc = DocxDocument(docx_path)
                week_numbers = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    # 匹配"第X周"或"第XX周"的模式
                    import re
                    match = re.search(r'第([一二三四五六七八九十百]+)周', text)
                    if match:
                        week_cn = match.group(1)
                        # 转换中文数字为阿拉伯数字
                        week_num = self._chinese_to_number(week_cn)
                        if week_num and week_num not in week_numbers:
                            week_numbers.append(week_num)
                    # 也支持"第31周"这样的阿拉伯数字
                    match = re.search(r'第(\d+)周', text)
                    if match:
                        week_num = int(match.group(1))
                        if week_num not in week_numbers:
                            week_numbers.append(week_num)
                
                week_numbers = sorted(week_numbers)
                if week_numbers:
                    # print(f"    文档包含周数: {week_numbers}")
                    # 如果是第一个晨兴文档,记录起始周数
                    if self.first_week_number is None:
                        self.first_week_number = min(week_numbers)
                        # print(f"    记录起始周数: {self.first_week_number}")
                    # 计算周数到章节的映射偏移量
                    week_offset = self.first_week_number - 1
                    # print(f"    周数映射: 第{min(week_numbers)}周 对应 第{min(week_numbers) - week_offset}篇")
                else:
                    week_offset = 0
            except Exception as e:
                print(f"    ⚠ 识别周数失败,使用顺序编号: {e}")
                week_numbers = []
                week_offset = 0
            
            # 从docx的zip结构中提取图片
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                # 列出所有图片文件
                image_files = sorted([f for f in docx_zip.namelist() 
                                    if f.startswith('word/media/') and not f.endswith('/')])
                
                # print(f"    找到 {len(image_files)} 个图片文件 {f'({doc_identifier})' if doc_identifier else ''}")
                
                for idx, img_file in enumerate(image_files):
                    # 使用识别到的周数,如果有的话
                    if week_numbers and idx < len(week_numbers):
                        actual_week_num = week_numbers[idx]
                        # 映射到章节索引：实际周数 - 偏移量 = 章节号
                        chapter_num = actual_week_num - week_offset
                    else:
                        actual_week_num = idx + 1 + week_offset  # 回退到顺序编号
                        chapter_num = idx + 1
                    
                    if chapter_num > len(chapters) or chapter_num < 1:
                        # print(f"    ⚠ 第{actual_week_num}周对应第{chapter_num}篇超出范围(1-{len(chapters)}),跳过")
                        continue
                    
                    # 添加文档标识符以避免覆盖
                    suffix = f"_{doc_identifier}" if doc_identifier else ""
                    # 使用章节号命名图片文件
                    image_path = os.path.join(output_dir, f'hymn_{chapter_num}{suffix}.png')
                    
                    # 如果图片文件已存在，跳过提取（避免重复提取）
                    if os.path.exists(image_path):
                        pass  # print(f"    跳过第{actual_week_num}周图片（文件已存在）")
                        # 但仍然需要记录路径到chapter
                        if not chapters[chapter_num - 1].hymn_image:
                            chapters[chapter_num - 1].hymn_image = f'images/hymn_{chapter_num}{suffix}.png'
                        continue
                    
                    try:
                        # 读取图片数据
                        img_data = docx_zip.read(img_file)
                        img = Image.open(BytesIO(img_data))
                        
                        # 转换为PNG并保存
                        img.save(image_path, 'PNG')
                        # print(f"    ✓ 第{actual_week_num}周 -> 第{chapter_num}篇诗歌图片{suffix}")
                        
                        # 记录图片路径到对应章节
                        chapters[chapter_num - 1].hymn_image = f'images/hymn_{chapter_num}{suffix}.png'
                    except Exception as e:
                        print(f"    ⚠ 保存第{actual_week_num}周图片失败: {e}")
            
        except ImportError as e:
            print(f"    ⚠ 缺少PIL库，无法提取图片: pip install pillow")
        except Exception as e:
            print(f"    ⚠ 图片提取异常: {e}")
            pass
            
    def _should_merge_with_previous(self, prev_text: str, current_text: str) -> bool:
        """判断当前文本是否应该与前一行文本合并（跨页连接）
        
        Args:
            prev_text: 前一行文本
            current_text: 当前行文本
            
        Returns:
            bool: True if should merge, False otherwise
        """
        if not prev_text or not current_text:
            return False
            
        # 保存原始文本用于检查缩进
        original_current = current_text
        prev_text_clean = prev_text.strip()
        current_text_clean = current_text.strip()
        
        # 如果前一行以完整句子标点结尾，通常不需要合并
        if prev_text_clean.endswith(('。', '！', '？', '；', '：', '）', '"', '”')):
            return False
            
        # 如果当前行很短，可能是标题或独立内容，不合并
        if len(current_text_clean) < 4:
            return False
            
        # 如果当前行以明显的段落标记开头，不合并
        if re.match(r'^[壹贰叁肆伍陆柒捌玖拾]+[、\s]', current_text_clean):
            return False
        if re.match(r'^[一二三四五六七八九十]+[、\s]', current_text_clean):
            return False
        if re.match(r'^\d+[、.\s]', current_text_clean):
            return False
        if re.match(r'^[a-z][、\.\s]', current_text_clean):
            return False
        if current_text_clean.startswith(('第', '(', '（', '[', '【')):
            return False
            
        # 关键修复：检查当前行是否有缩进
        # 如果当前行有缩进，说明不是跨页连接，不需要合并
        has_indent = original_current.startswith((' ', '\t'))
        if has_indent:
            return False
        
        # 如果当前行顶格开始（无缩进），才考虑跨页连接
        # 前一行明显未完成（以连接词、逗号等结尾），需要合并
        if prev_text_clean.endswith(('，', '、', '—', '－', '和', '与', '及', '或')):
            return True
        
        # 前一行没有完整结束，很可能是跨页连接
        # 前一行没有完整结束，很可能是跨页连接
        # 前一行没有标点结尾且足够长 - 可能被截断，需要合并
        if len(prev_text_clean) > 3 and not prev_text_clean[-1] in '。！？：；）"""】』':
            return True
        # 当前行以小写字母或中文开头，且前一行较长
        if len(prev_text_clean) > 3 and (
            re.match(r'^[a-z\u4e00-\u9fff]', current_text_clean) and 
            not re.match(r'^[壹贰叁肆伍陆一二三四五六七八九十]', current_text_clean)
        ):
            return True
            
        # 如果前一行很短且当前行不像新段落开始
        if (len(prev_text_clean) < 25 and len(current_text_clean) > 15 and 
            not re.match(r'^[壹贰叁肆伍陆一二三四五六七八九十\d]', current_text_clean)):
            return True
            
        return False
    
    def _clean_title(self, text: str) -> str:
        """清理标题，去掉层级标记后的内容"""
        # 去掉开头的层级标记（壹、一、1、a等）
        text = re.sub(r'^[壹贰叁肆伍陆柒捌玖拾]+\s+', '', text)
        text = re.sub(r'^[一二三四五六七八九十]+\s+', '', text)
        text = re.sub(r'^\d+\s+', '', text)
        text = re.sub(r'^[a-z]\s+', '', text)
        return text.strip()
    
    def _is_valid_ministry_text(self, text: str) -> bool:
        """
        判断文本是否为有效的职事信息内容
        
        过滤掉：
        1. 只包含下划线的文本
        2. 只包含空白字符的文本
        3. 过短的无意义文本
        
        Args:
            text: 待验证的文本
        
        Returns:
            True 如果是有效的职事信息内容，否则 False
        """
        if not text or not text.strip():
            return False
        
        # 去掉空白后检查
        cleaned = text.strip()
        
        # 如果全是下划线，认为无效
        if re.match(r'^_+$', cleaned):
            return False
        
        # 如果主要是下划线（超过80%），认为无效
        underscore_count = cleaned.count('_')
        if len(cleaned) > 0 and underscore_count / len(cleaned) > 0.8:
            return False
        
        # 长度过短且没有中文字符，可能是无效内容
        if len(cleaned) < 3 and not re.search(r'[\u4e00-\u9fff]', cleaned):
            return False
        
        return True


def _build_scripture_map(sections, scripture_map, prefix=""):
    """
    递归构建纲目经文映射
    
    Args:
        sections: Content对象列表
        scripture_map: 经文映射字典 {level+title -> scripture}
        prefix: 层级前缀
    """
    for section in sections:
        # 使用 level + title 作为键（去除空格以提高匹配率）
        key = (section.level + section.title).replace(' ', '').replace('\u3000', '')
        if section.scripture:
            scripture_map[key] = section.scripture
        
        # 递归处理子节点
        if section.children:
            _build_scripture_map(section.children, scripture_map, prefix + section.level)


def _fill_scripture(sections, scripture_map):
    """
    递归填充晨兴纲目的经文数据
    
    Args:
        sections: Content对象列表（晨兴纲目）
        scripture_map: 经文映射字典
    """
    for section in sections:
        # 使用 level + title 作为键查找经文
        key = (section.level + section.title).replace(' ', '').replace('\u3000', '')
        if key in scripture_map and not section.scripture:
            section.scripture = scripture_map[key]
        
        # 递归处理子节点
        if section.children:
            _fill_scripture(section.children, scripture_map)


def parse_training_docs_improved(outline_path: str, listen_path: str, 
                                 morning_revival_path: Optional[str] = None,
                                 morning_revival_path2: Optional[str] = None,
                                 title: str = "", subtitle: str = "", 
                                 year: int = 2025, season: str = "",
                                 output_dir: str = "output") -> TrainingData:
    """
    改进的文档解析流程
    
    Args:
        outline_path: 纲目文档路径（经文.docx）
        listen_path: 听抄文档路径（听抄.docx）
        morning_revival_path: 晨兴文档路径1（晨兴.doc），可选
        morning_revival_path2: 晨兴文档路径2（晨兴2.doc），可选
        title: 训练标题（如果为空则从文档中自动提取）
        subtitle: 副标题（如果为空则从文档中自动提取）
        year: 年份
        season: 季节
        output_dir: 输出目录
    """
    parser = ImprovedParser(output_dir=output_dir)
    
    # 1. 先解析纲目文档，建立大纲结构（同时提取标题）
    print("  解析纲目结构（经文.docx）...")
    chapters = parser.parse_outline_doc(outline_path)
    
    # 如果没有提供标题，使用从文档中提取的标题
    if not title and parser.training_title:
        title = parser.training_title
        print(f"  ✓ 使用文档中的标题: {title}")
    
    if not subtitle and parser.training_subtitle:
        subtitle = parser.training_subtitle
        print(f"  ✓ 使用文档中的副标题: {subtitle}")
    
    # 2. 再解析听抄文档，填充详细内容
    print("  解析详细内容（听抄.docx）...")
    parser.parse_listen_doc(listen_path, chapters)
    
    # 3. 解析晨兴文档（如果存在）
    if morning_revival_path:
        print("  解析晨兴内容（晨兴.doc）...")
        parser.parse_morning_revival_doc(morning_revival_path, chapters)
    
    if morning_revival_path2:
        print("  解析晨兴内容2（晨兴2.doc）...")
        parser.parse_morning_revival_doc(morning_revival_path2, chapters)
    
    # 3.5 将纲目的经文数据同步到晨兴纲目中（共用同一份经文数据）
    print("  同步经文数据到晨兴纲目...")
    for chapter in chapters:
        if chapter.outline_sections and chapter.morning_revivals:
            # 构建纲目经文映射：level+title -> scripture
            scripture_map = {}
            _build_scripture_map(chapter.outline_sections, scripture_map)
            
            # 为每天的晨兴纲目填充经文
            for revival in chapter.morning_revivals:
                if revival.outline:
                    _fill_scripture(revival.outline, scripture_map)
    
    # 4. 读取应用版本号
    app_version = ""
    try:
        import json
        with open('app_config.json', 'r', encoding='utf-8') as f:
            app_config = json.load(f)
            app_version = app_config.get('version', '')
    except Exception as e:
        print(f"  ⚠ 读取版本号失败: {e}")
    
    # 5. 创建训练数据对象
    training_data = TrainingData(
        title=title,
        subtitle=subtitle,
        year=year,
        season=season,
        app_version=app_version
    )
    
    for chapter in chapters:
        training_data.add_chapter(chapter)
    
    return training_data
